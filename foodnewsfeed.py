import streamlit as st
import feedparser
import pandas as pd
from datetime import datetime, timedelta
import io
from docx import Document
from docx.shared import Inches
import base64
import schedule
import time
from pytz import timezone
from googletrans import Translator  # For translation
import textwrap  # For wrapping summaries

# Define your list of RSS feeds
rss_feeds = {
    "EU Legislation": "https://eur-lex.europa.eu/EN/display-feed.rss?myRssId=e1Wry6%2FANeUpe1f1%2BCToXcHk31CTyaJK",
    "EFSA": "https://www.efsa.europa.eu/en/all/rss",
    "EU Food Safety": "https://food.ec.europa.eu/node/2/rss_en",
    "Food Quality & Safety": "https://www.foodqualityandsafety.com/category/eupdate/feed/",
    "Food Safety News": "https://feeds.lexblog.com/foodsafetynews/mRcs",
    "Food Manufacture": "https://www.foodmanufacture.co.uk/Info/FoodManufacture-RSS",
    "Food Packaging Forum": "https://www.foodpackagingforum.org/news/feed/",
    "ANSES": "https://www.anses.fr/fr/flux-actualites.rss" 
}

# Function to parse all RSS feeds (memoized to update once daily)
@st.experimental_memo(ttl=86400)  # TTL (Time-To-Live) is 1 day in seconds
def parse_feeds():
    df = pd.DataFrame()
    for feed_name, feed_url in rss_feeds.items():
        parsed_feed = feedparser.parse(feed_url)
        for entry in parsed_feed.entries:
            df = df.append({
                "feed": feed_name,
                "title": entry.title,
                "link": entry.link,
                "summary": entry.summary,
                "published": datetime.fromtimestamp(entry.published_parsed),
                "image": entry.get("media_content", [None])[0].get("url", None) if "media_content" in entry else None
            }, ignore_index=True)
    return df

# Get the current time in Paris
paris_timezone = timezone('Europe/Paris')
current_time = datetime.now(paris_timezone)

# Calculate the next scheduled update time (8:00 AM Paris time)
next_update_time = current_time.replace(hour=8, minute=0, second=0, microsecond=0)
if current_time > next_update_time:
    next_update_time += timedelta(days=1)

# Schedule the feed update
schedule.every().day.at("08:00").do(parse_feeds).run()

# Main app logic
st.title("Food Safety News & Reviews")

# Sidebar for navigation
with st.sidebar:
    st.header("Navigation")
    selected_feed = st.selectbox("Select a Feed:", list(parse_feeds()["feed"].unique()))
    st.write(f"Last Update: {datetime.now(paris_timezone).strftime('%Y-%m-%d %H:%M:%S')}")
    selected_language = st.selectbox("Select Language:", ["English", "French", "Spanish"])

st.markdown("---")
st.header(f"{selected_feed} Articles")

# Filter by date
selected_date = st.date_input("Select a Date", datetime.now())
df_filtered = parse_feeds()[(parse_feeds()["feed"] == selected_feed) & (parse_feeds()["published"].dt.date == selected_date)]

# Display filtered articles
for index, row in df_filtered.iterrows():
    st.write(f"**{row['title']}**")
    if row["image"]:
        st.image(row["image"], caption="Image from the article")

    # Translate summary
    translator = Translator()
    if selected_language == "French":
        translated_summary = translator.translate(row["summary"], dest="fr").text
    elif selected_language == "Spanish":
        translated_summary = translator.translate(row["summary"], dest="es").text
    else:
        translated_summary = row["summary"]

    # Wrap the summary for readability
    wrapped_summary = textwrap.wrap(translated_summary, width=100)
    for line in wrapped_summary:
        st.write(line)

    st.write(f"[Read More]({row['link']})")

    # Allow users to add articles to review
    review_button = st.button("Add to Review", key=f"review_{index}")
    if review_button:
        st.session_state["review_articles"] = st.session_state.get("review_articles", [])
        st.session_state["review_articles"].append(row)
        st.success(f"Article added to review!")

# Display review section
st.markdown("---")
st.header("Your Review")

# Display selected articles for review
if "review_articles" in st.session_state:
    st.write("Selected Articles:")
    for i, article in enumerate(st.session_state.get("review_articles", [])):
        st.write(f"**{i+1}. {article['title']}**")
        st.write(f"[Read More]({article['link']})")

    # Add review text area
    st.markdown("---")
    review_text = st.text_area("Add your review here:", height=150)

    # Download/Email Options
    st.markdown("---")
    download_format = st.selectbox("Download Format:", ["PDF", "Word", "Email"])
    if download_format == "PDF":
        # Convert review to PDF
        output = io.BytesIO()
        pdf_document = Document()
        pdf_document.add_paragraph("Food Safety News Review")
        pdf_document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        pdf_document.add_paragraph(f"Selected Articles:")
        for article in st.session_state.get("review_articles", []):
            pdf_document.add_paragraph(f"- {article['title']}")
            pdf_document.add_paragraph(f"{article['summary'][:200]}...")
        pdf_document.add_paragraph(review_text)
        pdf_document.save(output)
        output.seek(0)
        st.download_button(
            label="Download Review as PDF",
            data=output,
            file_name="food_safety_review.pdf",
            mime="application/pdf"
        )
    elif download_format == "Word":
        # Convert review to Word
        doc = Document()
        doc.add_paragraph("Food Safety News Review")
        doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
        doc.add_paragraph(f"Selected Articles:")
        for article in st.session_state.get("review_articles", []):
            doc.add_paragraph(f"- {article['title']}")
            doc.add_paragraph(f"{article['summary'][:200]}...")
        doc.add_paragraph(review_text)
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        st.download_button(
            label="Download Review as Word",
            data=output,
            file_name="food_safety_review.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    elif download_format == "Email":
        # (Easiest way for email)
        st.markdown(
            f"""
            <a href="mailto:?subject=Food Safety News Review&body={review_text}">Send Review via Email</a>
            """,
            unsafe_allow_html=True
        )
